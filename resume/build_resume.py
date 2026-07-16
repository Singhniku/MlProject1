from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/nikitasingh/Downloads/NikitaSingh_Resume_Updated.docx"
FONT = "Calibri"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.45)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.5)
sec.right_margin = Inches(0.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)

RIGHT_TAB = sec.page_width - sec.left_margin - sec.right_margin

def para(align=None, before=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p

def run(p, text, bold=False, italic=False, size=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return r

def bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "444444")
    pbdr.append(bottom)
    pPr.append(pbdr)

def header(text):
    p = para(before=9, after=3)
    run(p, text, bold=True, size=11)
    bottom_border(p)

def two_col(left, right, bold=False, italic=False, before=None, after=1):
    p = para(before=before, after=after)
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    run(p, left, bold=bold, italic=italic)
    run(p, "\t" + right, bold=bold, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.22)
    if bold_prefix:
        run(p, bold_prefix, bold=True)
    run(p, text)
    return p

# ---- Header ----
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
run(p, "Nikita Singh", bold=True, size=17)
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
run(p, "Gurugram, Haryana · nikitasingh18dec@gmail.com · 8318721284 · linkedin.com/in/nikitasingh98", size=10)

# ---- Experience ----
header("EXPERIENCE")

two_col("American Express", "Gurugram, Haryana", bold=True, before=4)
two_col("Software Engineer", "October 2024 – Present", italic=True)
bullet("Engineered an agentic AI assistant for the Caselite case-management platform using Llama-based LLMs, automating end-to-end case creation and case-processing workflows — reducing average case-handling time for Customer Care Professionals (CCPs) and establishing a scalable foundation for AI-driven servicing.")
bullet("Automated banking case creation in CLIC by onboarding the KYC Refresh Report case type with automated success/failure notifications and failure-reason reporting — eliminating manual case-creation effort, accelerating issue triage, and improving operational efficiency.")
bullet("Re-architected the demographic-fetching service in ACE using parallel processing with optimized pagination and sorting for high-volume linked-account scenarios (80+ accounts) — significantly reducing API latency, increasing throughput and scalability, and improving CCP user experience.")
bullet("Spearheaded the JDK 21 migration across the SCRA platform and 4 dependent repositories, modernizing the technology stack, improving runtime performance, and lowering long-term maintenance cost and technical debt.")
bullet("Integrated C360 Customer and Account REST APIs to enrich Linked Account data with Open Date, Military Lending Act (MLA), and State APR regulatory indicators, and optimized API orchestration to eliminate redundant GAR API calls — strengthening data accuracy, regulatory compliance, and response times.")
bullet("Designed and delivered configurable, template-driven communication frameworks with reason-code governance for SCRA benefit enrollment, banking dispute resolution, and Caselite, and automated acknowledgment emails post case creation — standardizing card-member outreach and reducing manual CCP workload.")
bullet("Built a centralized OneData UI abstracting and securing database CRUD operations — eliminating direct manual database access, mitigating credential-exposure risk, and laying the groundwork for role-based access control (RBAC).")
bullet("Delivered an Advanced Search capability (name, email, phone) for banking cases and onboarded new SCRA workbaskets and case types (KYC Refresh, BSA Attachment) — accelerating case discovery and reducing turnaround time for benefit-enrollment requests.")
bullet("Owned end-to-end resolution of critical production incidents for CBR cases, independently root-causing and fixing a correlation-ID mapping defect in ACE API integration; stabilized E2/E3 environments and executed zero-defect SCRA Phase-2 deployments across E3 and E3-SL (PIV), strengthening platform reliability and release quality.")
bullet("Remediated 100% of critical and high Twistlock security vulnerabilities in caselite-automation, contributed to the Jenkins-to-GitHub Actions CI/CD migration, and mentored new engineers through knowledge-transfer sessions on system design and delivery practices.")

two_col("Jubilant Foodworks Ltd.", "Gurugram, Haryana", bold=True, before=5)
two_col("Software Engineer", "August 2022 – September 2024", italic=True)
bullet("Led the design discussions and development of a high-availability nextgen order management system microservices architecture using SOLID principles and maximum code coverage, resulting in a 30% increase in system performance and reliability.")
bullet("Implemented MongoDB as the principal database solution, enhancing data storage and retrieval efficiency for a high-volume application.")
bullet("Employed Agile methodologies to implement Order History, Promo Integration in OMS, Advance Order, and Cancel Order, and continuously improved the functionality and performance of the Order Tracker and OMS Service.")
bullet("Built real-time data streaming pipelines using Kafka for Bulk Promo Rollback, capturing menu-sync events and order status updates for real-time analysis and faster business decision-making.")
bullet("Implemented Quartz Scheduler to identify and manage canceled orders — loyalty and promo rollback, refunds, and retries of failed advance orders for store delivery.")
bullet("Collaborated closely with 5+ cross-functional teams to gather requirements, define system architecture, and ensure alignment with business objectives.")

two_col("IndiaMart InterMesh Ltd.", "Noida, Uttar Pradesh", bold=True, before=5)
two_col("Software Engineer", "July 2021 – July 2022", italic=True)
bullet("Developed RESTful APIs in Spring Boot for reply mailers and their consumer files, optimizing communication processes and enhancing system efficiency.")
bullet("Created dynamic HTML email templates during the migration of consumer files to a centralized service, ensuring consistency across platforms and devices.")
bullet("Implemented CI/CD pipelines to automate the deployment process, streamlining development workflows.")

doc.add_page_break()
two_col("IndiaMart InterMesh Ltd.", "Noida, Uttar Pradesh", bold=True)
two_col("Software Engineer – Intern", "April 2021 – June 2021", italic=True)
bullet("Implemented RabbitMQ and Java to include reply snippets on the buyer side of SMS and personalised reply mailers.")

# ---- Skills ----
header("SKILLS")
for label, items in [
    ("Languages & Frameworks: ", "Java 21, Python, Spring Boot, Hibernate, REST APIs, Data Structures & Algorithms"),
    ("AI & Automation: ", "Agentic AI, LangChain, Hugging Face, Llama, Prompt Engineering"),
    ("Data & Messaging: ", "MongoDB, SQL, PostgreSQL, Kafka, RabbitMQ"),
    ("Platforms & DevOps: ", "Kubernetes, AWS, Git, GitHub Actions, Jenkins, CI/CD, BPMN"),
]:
    p = para(after=1.5)
    run(p, label, bold=True)
    run(p, items)

# ---- Certifications ----
header("CERTIFICATIONS")
bullet("Build with AI: Autonomous Agents with LangChain and Hugging Face")

# ---- Education ----
header("EDUCATION")
two_col("Harcourt Butler Technical University (H.B.T.U)", "Kanpur, Uttar Pradesh", bold=True)
two_col("Bachelor of Technology, Computer Science & Engineering · GPA: 8.215", "August 2017 – June 2021", italic=True)
two_col("Central Academy", "Basti, Uttar Pradesh", bold=True)
two_col("Intermediate · 92.20%", "March 2016 – May 2017", italic=True)
two_col("St. Basil's School", "Basti, Uttar Pradesh", bold=True)
two_col("High School · 89.33%", "March 2014 – May 2015", italic=True)

# ---- Leadership ----
header("LEADERSHIP EXPERIENCE")
two_col("Events Head", "HBTU · March 2019", bold=True)
p = para(after=1)
run(p, "Spearheaded a team of 60 as Events Head & Stage Management Head and successfully organized ADHYAAY-19, the Techno-Cultural Fest of HBTU.")

doc.save(OUT)
print("saved", OUT)
